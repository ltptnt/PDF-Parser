from Parsing import parse
from Doctor import Doctor
from Patient import Patient
from docx import Document, table, text, styles
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from typing import Callable


# CHANGE THIS LINE TO THE FILE YOU WANT TO PARSE
doctor = parse("example1.pdf")



variable_names = [
    "patient_full_name",
    "patient_full_name",
    "patient_dob_age",
    "patient_home_address",
    "patient_medi_number",
    "community_pharmacist",
    "patient_current_conditions",
    "doctor_full_name",
    "doctor_provider_num",
    "doctor_pref_contact",
    "request_time",
    "doctor_reason_for_referral",
    "medication_1",
    "med_dosage_1",
    "medication_2",
    "med_dosage_2",
    "patient_height",
    "patient_weight",
    "patient_blood_pressure",
    "patient_creatine",
    "patient_CrCl",
]


BLANK_VALUE = ""



def create_document(doctor: Doctor):
    document = Document("Blank.docx")

    # Build the document
    document = replace_text(doctor, document)
    document = replace_table_text(doctor, document)
    return document

def save_document(document: Document, doctor: Doctor): # type: ignore
    document.save(f"DMMR REPORT - {doctor.get_patient().get_name()}.docx") # type: ignore


def replace_text(doctor: Doctor, document: Document): # type: ignore
    # Replace the text in the document with the doctor's details
    is_next = False
    paragraphs = document.paragraphs
    for paragraph in paragraphs:
        text = paragraph.text
        inline = paragraph.runs
        for word in range(len(inline)):
            i = 0
            while i < len(variable_names):
                variable_name = variable_names[i]
                if variable_name in inline[word].text:
                    try:
                        if "doctor" in variable_name:
                            value = getattr(doctor, variable_name.split("_", 1)[1])
                        elif "patient" in variable_name:
                            value = getattr(doctor.get_patient(), variable_name.split("_", 1)[1])
                        else:
                            value = getattr(doctor.get_patient(), variable_name)
                        if value is None:
                            value = BLANK_VALUE
                        if type(value) == list:
                            # If the value is a list, then we need to format it with a new line for each item and a tab for each item
                            value = "\n".join(value)

                        inline[word].text = inline[word].text.replace(variable_name, value)
                        variable_names.remove(variable_name)
                        break
                    except AttributeError:
                        break
                i += 1
    return document


def replace_table_text(doctor: Doctor, document: Document): # type: ignore
    # Replace the text in the tables with the doctor's details
    conditions_table = document.tables[0]
    # Add the patient's current conditions to the table 
    for medication, dosage in doctor.get_patient().get_medications().items(): # type: ignore
        new_row = conditions_table.add_row()
        new_row.cells[0].text = medication
        new_row.cells[1].text = dosage
    
    return document


document = create_document(doctor)
save_document(document, doctor)




'''
document = Document("Blank.docx")
paragraphs = document.paragraphs

is_next = False
for paragraph in paragraphs:
    if "reason_for_referral" in paragraph.text:
        is_next = True
        continue
        
    if is_next:
        print(paragraph.text)
        inline = paragraph.runs
        for i in range(len(inline)):
            print(inline[i].text)
        break
'''