import pdfplumber
from datetime import date
from Doctor import Doctor
from Patient import Patient
from docx import Document


# Read in the two example files
def parse(file_path: str) -> Doctor:
    pdf = pdfplumber.open(file_path)
    # 0 = no template, 1 = template 1, 2 = template 2
    page = pdf.pages[0]
    text = page.extract_text()
    lines = text.split("\n")
    if "COMMUNITY PHARMACY DETAILS" in lines[1]:
        try:
            return template_1(pdf)
        except ValueError:
            raise ValueError("Invalid Format, try template 2")
    elif "ACCREDITED COMMUNITY PHARMACIST" in lines[1]:
        try:
            return template_2(pdf)
        except ValueError:
            raise ValueError("Invalid Format, try template 1")
    else:
        raise ValueError("File could not be parsed")



def template_1(pdf) -> Doctor:
    # For the first template, we'll extract the text from the first page
    # and then split it into lines
    page = pdf.pages[0]
    text = page.extract_text()
    lines = text.split("\n")

    new_doctor = Doctor()
    new_patient = Patient()
    pharmacist_name = None

    for line in lines:
        # Get doctor name, patient name
        if "Name:" in line and new_doctor.get_name() is None:
            new_doctor.set_name(line.split("Name:")[1].strip())
            continue
        elif "Name:" in line and new_patient.get_accredited_pharmacist() is None:
            pharmacist_name = line.split("Name:")[1].strip()
            new_patient.set_accredited_pharmacist(pharmacist_name)
            continue

        elif "Name:" in line and new_patient.get_name() is None:
            new_patient.set_name(line.split("Name:")[1].strip())
            continue

        # Get patient phone number, doctor phone number
        if "Phone:" in line and new_doctor.get_contact_number() is None:
            new_doctor.set_contact_number(line.split("Phone:")[1].strip())
            continue
        elif "Phone:" in line and new_patient.get_phone_number is None:
            new_patient.set_phone_number(line.split("Phone:")[1].strip())
            continue

        # Get Doctor address, patient address
        if "Address:" in line and new_doctor.get_address() is None:
            new_doctor.set_address(line.split("Address:")[1].strip())
            continue
        elif "Address:" in line and new_patient.get_address() is None:
            new_patient.set_address(line.split("Address:")[1].strip())
            continue

        # Get patient details
        if "D.O.B:" in line and new_patient.get_dob() is None:
            DOB = line.split("D.O.B:")[1].strip()
            # Convert DOB to datetime object
            date_of_birth = date(
                int(DOB.split("/")[2]), int(DOB.split("/")[1]), int(DOB.split("/")[0])
            )
            new_patient.set_dob(date_of_birth)

        if "Medicare No:" in line and new_patient.get_medicare() is None:
            new_patient.set_medicare(line.split("Medicare No:")[1].strip())
            continue

        if (
            "Community Pharmacy:" in line
            and new_patient.get_community_pharmacist() is None
        ):
            new_patient.set_community_pharmacist(
                line.split("Community Pharmacy:")[1].strip()
            )
            continue

        # Get doctor details
        if "Provider No:" in line and new_doctor.get_provider_number() is None:
            new_doctor.set_provider_number(line.split("Provider No:")[1].strip())
            continue

        if "Prescriber No:" in line and new_doctor.get_prescriber_number() is None:
            new_doctor.set_prescriber_number(line.split("Prescriber No:")[1].strip())
            continue

        if "Email:" in line and new_doctor.get_email() is None:
            new_doctor.set_email(line.split("Email:")[1].strip())
            continue
        # Get the patient's height, weight, blood pressure, and creatine
        if "Height:" in line and new_patient.get_height() is None:
            new_patient.set_height(line.split(" ")[1].strip())
            continue

        if "Weight:" in line and new_patient.get_weight() is None:
            new_patient.set_weight(line.split(" ")[1].strip())
            continue

        if "Blood Pressure:" in line and new_patient.get_blood_pressure() is None:
            new_patient.set_blood_pressure(line.split(" ")[2].strip())
            continue

        if "Creatinine:" in line and new_patient.get_creatine() is None:
            new_patient.set_creatine(line.split(" ")[1].strip())
            continue

    # Get the date an Currrent Conditions from the second page
    page = pdf.pages[1]
    text = page.extract_text()
    lines = text.split("\n")

    for line in lines:
        if "Date:" in line and new_doctor.get_request_time() is None:
            new_doctor.set_request_time(line.split(" ")[1].strip())
            continue
        if "Height:" in line and new_patient.get_height() is None:
            new_patient.set_height(line.split(" ")[1].strip())
            continue
        if "Weight:" in line and new_patient.get_weight() is None:
            new_patient.set_weight(line.split(" ")[1].strip())
            continue

    current_conditions = text.split("CURRENT CONDITIONS:")[1]
    current_conditions = current_conditions.split(
        "RELEVANT LABORATORY RESULTS AND BLOOD DRUG LEVELS (eg serum electrolytes, liver function test etc. as relevant)"
    )[0]
    current_conditions = current_conditions.split("\n")

    conditions = []
    for condition in current_conditions:
        if condition != "":
            conditions.append(condition.strip())

    new_patient.set_current_conditions(conditions)

    # Get the medications from the 2nd page
    # Crop the page to only include the left side of the page
    page = pdf.pages[1]
    # change the margin to be 0
    page = page.within_bbox((0, 0, page.width, page.height))
    left_side = page.crop((0, 0, page.width / 2, page.height))
    right_side = page.crop((page.width / 2, 0, page.width, page.height))
    table_strat = {
        "vertical_strategy": "text",
        "horizontal_strategy": "text",
        "text_x_tolerance": 1,
        "text_y_tolerance": 1,
    }
    l_tables = left_side.extract_table(table_settings=table_strat)
    r_text = right_side.extract_text()
    dosages = r_text.split("\n")

    medications = []
    for table in l_tables:
        medication = "".join(table)
        medications.append(medication)
    # Split the list to only the first part of the list excluding the title
    # which is the medications
    medications = medications[4: medications.index("CURRENT CONDITIONS:") - 1]

    dosages = dosages[
        0: dosages.index(
            "S (eg serum electrolytes, liver function test etc. as relevant)"
        )
    ]

    # covert the dosages to a dictionary.
    # If there is a medication with no value then the values for the next two
    # medications will be the value for the previous medication
    medication_dict = {}
    multi_line_adjustment = 0
    for i in range(len(medications)):
        if medications[i] == "":
            multi_line_medication = (
                dosages[i - 1] + dosages[i] + dosages[i + 1] + dosages[i + 2]
            )
            multi_line_adjustment += 2
            medication_dict[medications[i - 1]] = multi_line_medication
        else:
            medication_dict[medications[i]] = dosages[i +
                                                      multi_line_adjustment]

    # Add the medications to the patient
    new_patient.set_medications(medication_dict)

    # Set the doctor and patient
    new_doctor.set_patient(new_patient)
    return new_doctor


def template_2(pdf) -> Doctor:
    page = pdf.pages[0]
    text = page.extract_text()
    lines = text.split("\n")

    new_doctor = Doctor()
    new_patient = Patient()
    pharmacist_name = None

    BP_readings = {}
    Height_readings = {}
    Weight_readings = {}

    for line in lines:
        # Get accredited pharmacist name
        if "ACCREDITED COMMUNITY PHARMACIST’S DETAILS:" in line and new_patient.get_accredited_pharmacist() is None:
            pharmacist_name = line.split("ACCREDITED COMMUNITY PHARMACIST’S DETAILS:")[1].strip()
            new_patient.set_accredited_pharmacist(pharmacist_name)
            continue
        # Get doctor name, patient name
        if "Patient Name:" in line and new_patient.get_name() is None:
            # get text between "Patient Name:" and Drs Name
            new_patient.set_name(line.split("Patient Name:")[1].split("Drs. Name")[0].strip())
            new_doctor.set_name(line.split("Drs. Name:")[1].strip())
            continue
        # Get Patient Address
        if "Patient Address:" in line and new_patient.get_address() is None:
            new_patient.set_address(line.split("Patient Address:")[1].strip())
            continue
        # Get patient phone number and provider number
        if "Patient Phone:" in line and new_patient.get_phone_number() is None:
            new_patient.set_phone_number(line.split("Patient Phone:")[1].split("Provider No:")[0].strip())
            new_doctor.set_provider_number(line.split("Provider No:")[1].strip())
            continue
        # Get DOB and Prescriber No
        if "DOB:" in line and new_patient.get_dob() is None:
            DOB = line.split("DOB:")[1].split("Prescriber No:")[0].strip()
            new_patient.set_dob(DOB)
            new_doctor.set_prescriber_number(line.split("Prescriber No:")[1].strip())
            continue
        # Get Medicare No and Phone
        if "Medicare No:" in line and new_patient.get_medicare() is None:
            new_patient.set_medicare(line.split("Medicare No:")[1].split("Phone:")[0].strip())
            new_doctor.set_contact_number(line.split("Phone:")[1].strip())
            continue
        
        if "BP" in line:
            # Format date "BP" reading "Sitting"
            bp_measure = line.split(" ")[0]
            measurement_date = date(int(bp_measure.split("/")[2]), int(bp_measure.split("/")[1]), int(bp_measure.split("/")[0]))
            BP_readings[measurement_date] = line.split(" ")[2]
            continue
        if "Height" in line:
            # Format date "Height" reading
            height_measure = line.split(" ")[0]
            measurement_date = date(int(height_measure.split("/")[2]), int(height_measure.split("/")[1]), int(height_measure.split("/")[0]))
            Height_readings[measurement_date] = line.split(" ")[2]
            continue
        if "Weight" in line:
            # Format date "Weight" reading
            weight_measure = line.split(" ")[0]
            measurement_date = date(int(weight_measure.split("/")[2]), int(weight_measure.split("/")[1]), int(weight_measure.split("/")[0]))
            Weight_readings[measurement_date] = line.split(" ")[2]
            continue

    # Get the most recent reading
    new_patient.set_blood_pressure(BP_readings[max(BP_readings)])
    new_patient.set_height(Height_readings[max(Height_readings)] + "cm")
    new_patient.set_weight(Weight_readings[max(Weight_readings)] + "kg")

    # Get the medications from the 2nd page
    # get the left side of the page
    page = pdf.pages[1]
    # change the margin to be 0
    page = page.within_bbox((0, 0, page.width, page.height))
    left_side = page.crop((0, 0, page.width / 2, page.height))
    
    # extract text from left side
    l_text = left_side.extract_text()
    l_text = l_text.split("\n")
    # remove all values before and including "Current Medications"
    l_text = l_text[l_text.index("CURRENT MEDICATION:") + 1:]
    # remove all values after and including "Current Conditions"
    l_text = l_text[:l_text.index("CURRENT CONDITIONS:")]
    
    lines = page.extract_text().split("\n")
    # remove all values before and including "CURRENT MEDICATION:"
    lines = lines[lines.index("CURRENT MEDICATION:") + 1:]
    # remove all values after and including "CURRENT CONDITIONS:"
    lines = lines[:lines.index("CURRENT CONDITIONS:")]

    # split lines into medications and dosages using the l_text medication list
    medications = []
    dosages = []
    for line in lines:
        for medication in l_text:
            if medication in line:
                medications.append(medication)
                dosages.append(line.split(medication)[1].strip())
                # if the line is not the last line and the last character is not a period
                # then the dosage is a multi-line dosage
                if line != lines[-1] and line[-1] != ".":
                    dosages[-1] += " " + lines[lines.index(line) + 1]
                break

    # convert the medications and dosages to a dictionary
    medication_dict = dict(zip(medications, dosages))
    # Add the medications to the patient

    # Set the doctor and patient
    new_patient.set_medications(medication_dict)

    # Get current conditions from the 2nd page and 3rd page
    page = pdf.pages[1]
    text = page.extract_text()
    lines = text.split("\n")
    for line in lines:
        if "CURRENT CONDITIONS:" in line:
            lines = lines[lines.index(line) + 1:]
            break
    current_conditions = []
    current_conditions += lines

    # Get the date and rest of the current conditions from the 3rd page
    page = pdf.pages[2]
    # adjust the margin to be 0
    page = page.within_bbox((0, 0, page.width, page.height))
    text = page.extract_text()
    lines = text.split("\n")
    # if the first line contains "//", remove every second character from the line
    if "//" in lines[0]:
        fixed = ""
        for word in lines[0].split(" "):
            word = word[::2]
            fixed += word + " "
        lines[0] = fixed
    # split after "CURRENT CONDITIONS:" and before "RELEVANT LABORATORY RESULTS AND BLOOD DRUG LEVELS"
    for line in lines:
        if "CURRENT CONDITIONS:" in line:
            lines = lines[lines.index(line) + 1:]
            break

    for line in lines:
        if "RELEVANT LABORATORY" in line:
            lines = lines[:lines.index(line)]
            continue
 
    current_conditions += (lines)
    new_patient.set_current_conditions(current_conditions)

    # Get date from the last page
    page = pdf.pages[-1]
    text = page.extract_text()
    lines = text.split("\n")
    for line in lines:
        if "Date:" in line:
            new_doctor.set_request_time(line.split(" ")[1].strip())
            break

    new_doctor.set_patient(new_patient)

    return new_doctor


def create_document(doctor: Doctor):
    document = Document("resources/Blank.docx")
    if doctor is None:
        return document

    variable_names = [
        "patient_full_name",
        "patient_full_name",
        "patient_dob_age",
        "patient_home_address",
        "patient_medi_number",
        "community_pharmacist",
        "patient_current_conditions",
        "doctor_full_name",
        "doctor_provider_num",
        "doctor_pref_contact",
        "doctor_request_time",
        "doctor_reason_for_referral",
        "medication_1",
        "med_dosage_1",
        "medication_2",
        "med_dosage_2",
        "patient_height",
        "patient_weight",
        "patient_blood_pressure",
        "patient_creatine",
        "patient_CrCl",
    ]

    # Replace the text in the document with the doctor's details
    paragraphs = document.paragraphs
    for paragraph in paragraphs:
        inline = paragraph.runs
        for word in range(len(inline)):
            i = 0
            while i < len(variable_names):
                variable_name = variable_names[i]
                if variable_name in inline[word].text:
                    try:
                        if "doctor" in variable_name:
                            value = getattr(doctor,
                                            variable_name.split("_", 1)[1])
                        elif "patient" in variable_name:
                            value = getattr(
                                doctor.get_patient(),
                                variable_name.split("_", 1)[1]
                            )
                        else:
                            value = getattr(doctor.get_patient(),
                                            variable_name)
                        if value is None:
                            value = ""
                        if type(value) is list:
                            # If the value is a list, then we need to format it
                            # with a new line and a tab for each item
                            value = "\n".join(value)

                        inline[word].text = inline[word].text.replace(
                            variable_name, value
                        )
                        variable_names.remove(variable_name)
                        break
                    except AttributeError:
                        break
                i += 1

    # Replace the text in the tables with the doctor's details
    conditions_table = document.tables[0]
    # Add the patient's current conditions to the table
    for medication, dosage in doctor.get_patient().get_medications().items():
        new_row = conditions_table.add_row()
        new_row.cells[0].text = medication
        new_row.cells[1].text = dosage

    return document